from sys import argv
from docx import Document
from enum import Enum
from database import *
import datetime

def xstr(s):
    if s is None:
        return ''
    return str(s)

class Level(Enum):
    CLUB = "Cl"
    COUNTY = "Co"
    DISTRICT = "D"
    STATE = "S"

class LeadershipRole(Enum):
    ELECTED = "E"
    APPOINTED = "A"
    VOLUNTEER = "V"
    PROMOTIONAL = "P"

class ServiceRole(Enum):
    YOURSELF = "Y"
    MEMBER = "M"
    PRIMARY = "P"

class JournalWriter:

    def __init__(self, document_name):
        self.document = Document(document_name)
        self.document_name = document_name


    def create_project_table(self, project_name, experiences):
        project_title = self.document.add_paragraph()
        project_title.text = project_name

        project_table = self.document.add_table(len(experiences) + 1, 4)

        keys = ["Year", "Project Activity", "Hours", "Importance to You"]
        j = 0
        for cell in project_table.rows[0].cells:
            run = cell.paragraphs[0].add_run(keys[j])
            run.bold = True
            j += 1
        i = 1
        for experience in experiences:

            project_table.cell(i, 0).text = experience.year
            project_table.cell(i, 1).text = experience.activity
            project_table.cell(i, 2).text = experience.hours
            project_table.cell(i, 3).text = experience.importance
            i += 1
        self.document.save(self.document_name)



class RecordbookWriter:

    def __init__(self, document_name):
        self.document = Document(document_name)
        self.document_name = document_name
        tables = self.document.tables

        self.name = tables[0].cell(0, 1)
        self.county = tables[0].cell(0, 3)
        self.district = tables[0].cell(0, 5)
        self.division = tables[0].cell(1, 1)
        self.category = tables[0].cell(1, 3)

        self.leadership_form = tables[6].rows[3:]
        self.service_form = tables[8].rows[3:]
        self.award_form = tables[12].rows[3:]
        self.career_form = tables[16].rows[6:]

    def row_empty(self, row):
        i = 0
        for cell in row.cells:
            if cell.text and i != 0:
                return False
            i += 1
        return True

    def get_empty_row(self, rows):
        curr_row = None
        for row in rows:
            if self.row_empty(row):
                curr_row = row
                break
        return curr_row


    def fill_info(self, name="", county="", district="", division="", category=""):
        self.name.text = xstr(name)
        self.county.text = xstr(county)
        self.district.text = xstr(district)
        self.division.text = xstr(division)
        self.category.text = xstr(category)
        self.document.save(self.document_name)

    def append_leadership(self, activity, role, level, year=str(datetime.datetime.today().year), duties=""):


        curr_row = self.get_empty_row(self.leadership_form)
        if not curr_row:
            curr_row = self.document.tables[6].add_row()
        row_cells = curr_row.cells

        row_cells[1].text = xstr(year)
        row_cells[2].text = xstr(activity)
        row_cells[3].text = xstr(role)
        row_cells[4].text = xstr(level)
        row_cells[5].text = xstr(duties)
        self.document.save(self.document_name)

    def append_service(self, role, activity, year=str(datetime.datetime.today().year), impact=""):

        curr_row = self.get_empty_row(self.service_form)
        if not curr_row:
            curr_row = self.document.tables[8].add_row()
        row_cells = curr_row.cells

        row_cells[1].text = xstr(year)
        row_cells[2].text = xstr(role)
        row_cells[3].text = xstr(activity)
        row_cells[4].text = xstr(impact)
        self.document.save(self.document_name)

    def append_award(self, level, recognition, year=str(datetime.datetime.today().year), importance=""):

        curr_row = self.get_empty_row(self.award_form)
        if not curr_row:
            curr_row = self.document.tables[12].add_row()
        row_cells = curr_row.cells

        row_cells[1].text = xstr(year)
        row_cells[2].text = xstr(level)
        row_cells[3].text = xstr(recognition)
        row_cells[4].text = xstr(importance)
        self.document.save(self.document_name)

    def append_career(self, activity, year=str(datetime.datetime.today().year), importance=""):
        curr_row = self.get_empty_row(self.career_form)
        if not curr_row:
            curr_row = self.document.tables[16].add_row()
        row_cells = curr_row.cells

        row_cells[1].text = xstr(year)
        row_cells[2].text = xstr(activity)
        row_cells[3].text = xstr(importance)
        self.document.save(self.document_name)
